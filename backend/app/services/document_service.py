import asyncio
import os
import tempfile
import base64
from typing import List, Dict, Any, Optional, Union, Callable
from pathlib import Path
import aiofiles
import aiofiles.os
from docx import Document

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
import io
from app.services.ollama_service import ollama_service
from app.config.settings import settings

# OCR and image conversion dependencies
try:
    import pytesseract
except ImportError:
    pytesseract = None
try:
    from pdf2image import convert_from_path
except ImportError:
    convert_from_path = None
from PIL import Image


class ModelRegistry:
    """Registry of model capabilities and context windows"""

    # Known model context windows (in tokens)
    MODEL_CONTEXTS = {
        "tinyllama:1.1b": 2048,
        "llama3.2:1b": 2048,
        "llama3.2:3b": 4096,
        "llama3.2:8b": 8192,
        "phi3:mini": 8192,
        "llama3:latest": 8192,
        "codellama:7b-instruct": 8192,
        "deepseek-coder:6.7b": 8192,
        "dolphin-mistral:7b": 8192,
        "starcoder2:3b": 8192,
        "qwen2.5-coder:1.5b": 8192,
        "qwen2.5-coder:3b": 8192,
        "gemma:2b": 8192,
        # Vision models
        "llava:7b": 8192,
        "llava:13b": 8192,
        "llava:34b": 8192,
        "bakllava:7b": 8192,
        "llava-llama3.2:8b": 8192,
    }

    # Models known to be slower and need more aggressive chunking
    SLOW_MODELS = {
        "dolphin-mistral:7b",
        "deepseek-coder:6.7b",
        "codellama:7b-instruct",
    }

    # Small models that need special handling for document analysis
    SMALL_MODELS = {
        "tinyllama:1.1b",
        "llama3.2:1b",
    }

    # Vision models that support image analysis
    VISION_MODELS = {
        "llava:7b",
        "llava:13b",
        "llava:34b",
        "bakllava:7b",
        "llava-llama3.2:8b",
    }

    # Default context window for unknown models
    DEFAULT_CONTEXT = 4096

    @classmethod
    def get_context_length(cls, model_name: str) -> int:
        """Get context length for a model"""
        return cls.MODEL_CONTEXTS.get(model_name, cls.DEFAULT_CONTEXT)

    @classmethod
    def estimate_tokens(cls, text: str) -> int:
        """Rough estimate of tokens in text (4 chars per token is a reasonable approximation)"""
        return len(text) // 4

    @classmethod
    def should_chunk(cls, text: str, model_name: str, prompt_length: int = 500) -> bool:
        """Determine if text should be chunked based on model context"""
        context_length = cls.get_context_length(model_name)
        estimated_tokens = cls.estimate_tokens(text) + prompt_length

        # Leave some buffer for response (about 25% of context)
        available_tokens = int(context_length * 0.75)

        should_chunk_result = estimated_tokens > available_tokens
        print(
            f"[DEBUG] ModelRegistry.should_chunk: text_len={len(text)}, model={model_name}, context_length={context_length}, estimated_tokens={estimated_tokens}, available_tokens={available_tokens}, should_chunk={should_chunk_result}"
        )

        return should_chunk_result

    @classmethod
    def get_optimal_chunk_size(cls, model_name: str, prompt_length: int = 500) -> int:
        """Get optimal chunk size for a model"""
        context_length = cls.get_context_length(model_name)
        available_tokens = int(context_length * 0.75) - prompt_length

        # Convert tokens to characters (rough estimate)
        base_size = available_tokens * 4

        # For small models, use much smaller chunks to ensure they can process effectively
        if model_name in cls.SMALL_MODELS:
            return int(base_size * 0.4)  # 40% of normal size for small models
        # For slow models, use much smaller chunks
        elif model_name in cls.SLOW_MODELS:
            return int(base_size * 0.3)  # 30% of normal size for slow models

        return base_size

    @classmethod
    def is_small_model(cls, model_name: str) -> bool:
        """Check if a model is considered small and needs special handling"""
        return model_name in cls.SMALL_MODELS

    @classmethod
    def is_vision_model(cls, model_name: str) -> bool:
        """Check if a model supports vision/image analysis"""
        return model_name in cls.VISION_MODELS

    @classmethod
    async def update_model_info(cls, model_name: str) -> None:
        """Dynamically update model information from Ollama"""
        try:
            model_info = await ollama_service.get_model_info(model_name)
            if "model_info" in model_info:
                info = model_info["model_info"]
                if "llama.context_length" in info:
                    context_length = info["llama.context_length"]
                    cls.MODEL_CONTEXTS[model_name] = context_length
                    print(f"Updated context length for {model_name}: {context_length}")
        except Exception as e:
            print(f"Could not update model info for {model_name}: {e}")


class DocumentService:
    """Service for analyzing and processing various document types"""

    SUPPORTED_EXTENSIONS = {
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".pdf": "application/pdf",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".csv": "text/csv",
        ".json": "application/json",
        ".xml": "application/xml",
        ".html": "text/html",
        ".htm": "text/html",
    }

    def __init__(self):
        self.max_chunk_size = 4000  # Tokens per chunk (legacy, will be overridden)
        self.max_total_tokens = 32000  # Total tokens for analysis (legacy)

    @property
    def document_timeout(self) -> float:
        """Get document timeout from settings, with fallback to default"""
        return settings.document_timeout

    def calculate_timeout(
        self,
        text_length: int,
        is_chunked: bool = False,
        model_name: Optional[str] = None,
    ) -> float:
        """Calculate appropriate timeout based on document size and model"""
        base_timeout = self.document_timeout

        # For slow models, use shorter timeouts
        if model_name and model_name in ModelRegistry.SLOW_MODELS:
            base_timeout = base_timeout * 0.6  # 60% of normal timeout for slow models

        if is_chunked:
            # For chunked analysis, use shorter timeouts per chunk
            # For very large documents, use more reasonable timeouts
            if text_length > 20000:
                # Use more reasonable timeout for large documents, especially with better models
                if model_name and model_name in ["phi3:mini", "llama3:latest"]:
                    return min(
                        45.0, base_timeout / 3
                    )  # Better timeout for capable models
                else:
                    return min(
                        25.0, base_timeout / 5
                    )  # Still reasonable for other models
            elif model_name and model_name in ModelRegistry.SLOW_MODELS:
                return min(25.0, base_timeout / 5)  # Reasonable timeout for slow models
            else:
                return min(40.0, base_timeout / 3)  # Standard chunked timeout

        # For direct analysis (including final summaries), scale timeout based on document size
        # Be much more generous with timeouts for small documents to account for model loading time
        if text_length < 500:
            # Very small documents - use very generous timeout to account for model startup time
            return min(120.0, base_timeout * 1.2)  # 120% of base timeout, max 120s
        elif text_length < 1000:
            # Small documents - use generous timeout for model loading
            return min(90.0, base_timeout * 1.0)  # 100% of base timeout, max 90s
        elif text_length < 5000:
            # Medium documents - use generous timeout
            return min(120.0, base_timeout * 1.0)  # 100% of base timeout, max 120s
        elif text_length < 15000:
            # Large documents use full timeout
            return base_timeout
        elif text_length < 30000:
            # Very large documents - use generous timeout
            return min(180.0, base_timeout * 1.5)  # 150% of base timeout, max 3 minutes
        elif text_length < 50000:
            # Extremely large documents - use very generous timeout
            return min(300.0, base_timeout * 2.0)  # 200% of base timeout, max 5 minutes
        else:
            # Massive documents - use maximum reasonable timeout
            return min(
                600.0, base_timeout * 3.0
            )  # 300% of base timeout, max 10 minutes

    def get_model_aware_chunk_size(self, model_name: str) -> int:
        """Get chunk size optimized for the specific model"""
        return ModelRegistry.get_optimal_chunk_size(model_name)

    def should_chunk_documents(self, combined_text: str, model_name: str) -> bool:
        """Determine if documents should be chunked based on model capabilities and document size"""
        print(
            f"[DEBUG] should_chunk_documents called with text length: {len(combined_text)}, model: {model_name}"
        )

        # Check if it exceeds model context window
        context_chunk_needed = ModelRegistry.should_chunk(combined_text, model_name)
        print(f"[DEBUG] Context chunk needed: {context_chunk_needed}")

        if context_chunk_needed:
            print(f"[DEBUG] Returning True due to context window")
            return True

        # Also chunk if document is very large (>15k chars) to prevent timeouts
        # Large documents can cause timeouts even if they fit in context
        size_chunk_needed = len(combined_text) > 15000
        print(
            f"[DEBUG] Size chunk needed: {size_chunk_needed} (text length: {len(combined_text)})"
        )

        if size_chunk_needed:
            print(
                f"[DocumentService] Document is large ({len(combined_text)} chars), chunking to prevent timeouts"
            )
            print(f"[DEBUG] Returning True due to document size")
            return True

        print(f"[DEBUG] Returning False - no chunking needed")
        return False

    def chunk_text(
        self, text: str, model_name: str, chunk_overlap: int = 200
    ) -> List[str]:
        """Split text into chunks optimized for the model"""
        # Get optimal chunk size for the model
        optimal_size = self.get_model_aware_chunk_size(model_name)

        # For very large documents, use smaller chunks to prevent timeouts
        if len(text) > 20000:
            # Use much smaller chunks for large documents to prevent timeouts
            if model_name in ["phi3:mini", "llama3:latest"]:
                chunk_size = int(
                    optimal_size * 0.4
                )  # Use 40% of optimal size for large documents
                print(
                    f"[DocumentService] Using conservative chunks ({chunk_size} chars) for large document with {model_name} ({len(text)} chars)"
                )
            else:
                chunk_size = int(
                    optimal_size * 0.3
                )  # Use even smaller chunks for other models
                print(
                    f"[DocumentService] Using smaller chunks ({chunk_size} chars) for large document ({len(text)} chars)"
                )
        elif len(text) > 10000:
            # For medium-large documents, use moderate chunk sizes
            if model_name in ["phi3:mini", "llama3:latest"]:
                chunk_size = int(optimal_size * 0.6)  # Use 60% of optimal size
                print(
                    f"[DocumentService] Using moderate chunks ({chunk_size} chars) for medium document with {model_name} ({len(text)} chars)"
                )
            else:
                chunk_size = int(optimal_size * 0.5)
                print(
                    f"[DocumentService] Using moderate chunks ({chunk_size} chars) for medium document ({len(text)} chars)"
                )
        elif model_name in ModelRegistry.SLOW_MODELS:
            # For slow models, use even smaller chunks
            chunk_size = int(optimal_size * 0.5)
            print(
                f"[DocumentService] Using smaller chunks ({chunk_size} chars) for slow model {model_name} ({len(text)} chars)"
            )
        else:
            chunk_size = optimal_size

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # If this isn't the last chunk, try to break at a sentence boundary
            if end < len(text):
                # Look for sentence endings within the last 200 characters of the chunk
                search_start = max(start, end - 200)
                for i in range(end, search_start, -1):
                    if text[i] in ".!?\n":
                        end = i + 1
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            # Move start position, accounting for overlap
            start = end - chunk_overlap
            if start >= len(text):
                break

        print(f"[DocumentService] Split document into {len(chunks)} chunks")
        return chunks

    def is_supported_file(self, filename: str) -> bool:
        """Check if file type is supported"""
        ext = Path(filename).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS

    async def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """Save uploaded file to temporary location"""
        # Create temp directory if it doesn't exist
        temp_dir = Path(tempfile.gettempdir()) / "elara_documents"
        temp_dir.mkdir(exist_ok=True)

        # Generate unique filename
        file_path = temp_dir / f"{os.urandom(8).hex()}_{filename}"

        print(
            f"[DocumentService] Saving file: {filename}, content size: {len(file_content)} bytes"
        )

        async with aiofiles.open(file_path, "wb") as f:
            await f.write(file_content)

        # Verify file was saved correctly
        saved_size = os.path.getsize(file_path)
        print(f"[DocumentService] File saved: {file_path}, size: {saved_size} bytes")
        if saved_size != len(file_content):
            print(
                f"[DocumentService] WARNING: File size mismatch! Expected: {len(file_content)}, Actual: {saved_size}"
            )

        return str(file_path)

    async def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n\n".join(text_parts)
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX file: {str(e)}")

    async def extract_text_from_pdf_ocr(self, file_path: str) -> str:
        """Extract text from image-based PDF using OCR (pytesseract)"""
        if PyPDF2 is None:
            raise Exception(
                "PyPDF2 is not installed. Please install it with: pip install PyPDF2"
            )
        if pytesseract is None:
            raise Exception(
                "pytesseract is not installed. Please install it with: pip install pytesseract"
            )
        if convert_from_path is None:
            raise Exception(
                "pdf2image is not installed. Please install it with: pip install pdf2image"
            )
        try:
            images = convert_from_path(file_path)
            text_parts = []
            for page_num, pil_image in enumerate(images):
                try:
                    ocr_text = pytesseract.image_to_string(pil_image)
                    text_parts.append(f"Page {page_num + 1} (OCR):\n{ocr_text}")
                except Exception as e:
                    text_parts.append(f"Page {page_num + 1} (OCR) failed: {e}")
            # Always return a string
            return "\n\n".join(text_parts)
        except Exception as e:
            return f"OCR extraction failed: {str(e)}"

    async def extract_text_from_pdf(self, file_path: str) -> Union[str, Dict[str, str]]:
        """Extract text from PDF file"""
        if PyPDF2 is None:
            raise Exception(
                "PyPDF2 is not installed. Please install it with: pip install PyPDF2"
            )
        try:
            # Add file size debugging
            file_size = os.path.getsize(file_path)
            print(f"[DocumentService] PDF file size: {file_size} bytes")
            text_parts = []
            total_pages = 0
            pages_with_text = 0
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                print(f"[DocumentService] PDF has {total_pages} pages")
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(f"Page {page_num + 1}:\n{page_text}")
                            pages_with_text += 1
                            print(
                                f"[DocumentService] Page {page_num + 1} has {len(page_text)} characters"
                            )
                        else:
                            print(
                                f"[DocumentService] Page {page_num + 1} has no extractable text (may be image-based)"
                            )
                    except Exception as page_error:
                        print(
                            f"[DocumentService] Error extracting text from page {page_num + 1}: {page_error}"
                        )
            extracted_text = "\n\n".join(text_parts)
            print(
                f"[DocumentService] PDF extraction summary: {total_pages} total pages, {pages_with_text} pages with text, {len(extracted_text)} total characters"
            )
            if not extracted_text.strip():
                # Check if this might be an image-based PDF
                if total_pages > 0:
                    # Instead of raising, return a special dict
                    return {
                        "type": "image_based_pdf",
                        "message": "This PDF appears to be a scanned document. Would you like to extract text using OCR or analyze it as an image with the vision model?",
                        "file_path": file_path,
                    }
                else:
                    raise Exception("PDF appears to be empty or corrupted")
            return extracted_text
        except Exception as e:
            if "image-based" in str(e) or "no extractable text" in str(e):
                raise e  # Re-raise our specific error
            else:
                raise Exception(f"Failed to extract text from PDF file: {str(e)}")

    async def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                content = await f.read()
            return content
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, "r", encoding="latin-1") as f:
                    content = await f.read()
                return content
            except Exception as e:
                raise Exception(f"Failed to read text file with any encoding: {str(e)}")
        except Exception as e:
            raise Exception(f"Failed to extract text from text file: {str(e)}")

    async def extract_text_from_file(
        self, file_path: str, filename: str
    ) -> Union[str, Dict[str, str]]:
        """Extract text from file based on its extension"""
        ext = Path(filename).suffix.lower()

        if ext == ".docx":
            return await self.extract_text_from_docx(file_path)
        elif ext == ".pdf":
            return await self.extract_text_from_pdf(file_path)
        elif ext in [".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm"]:
            return await self.extract_text_from_txt(file_path)
        else:
            raise Exception(f"Unsupported file type: {ext}")

    async def _warm_up_model(self, model: str) -> None:
        """Warm up the model by sending a simple request to load it into memory"""
        try:
            print(f"[DocumentService] Warming up model: {model}")
            warm_up_prompt = "Hello, this is a warm-up request."
            warm_up_timeout = 30.0  # Short timeout for warm-up

            await ollama_service.query_ollama(warm_up_prompt, warm_up_timeout, model)
            print(f"[DocumentService] Model {model} warmed up successfully")
        except Exception as e:
            print(f"[DocumentService] Model warm-up failed for {model}: {e}")
            # Don't raise the exception - warm-up failure shouldn't stop the main process
            pass

    async def analyze_documents(
        self,
        files: List[Dict[str, Union[str, bytes]]],
        prompt: str,
        model: Optional[str] = None,
        progress_callback: Optional[Callable[[str, Union[int, float]], None]] = None,
        stop_event: Optional[asyncio.Event] = None,
    ) -> Dict[str, Any]:
        if model is None:
            model = settings.DOCUMENT_ANALYSIS_MODEL

        # Try to update model info if not in registry
        if model not in ModelRegistry.MODEL_CONTEXTS:
            await ModelRegistry.update_model_info(model)

        print(
            f"[DocumentService] Starting analysis: model={model}, files={len(files)}, prompt='{prompt[:50]}...', settings_timeout={self.document_timeout}"
        )

        # Check if model is available (basic check)
        try:
            import httpx

            async with httpx.AsyncClient(timeout=5.0) as client:
                response = await client.get(f"http://localhost:11434/api/tags")
                if response.status_code != 200:
                    print(
                        f"[DocumentService] WARNING: Ollama service may not be responding properly (status: {response.status_code})"
                    )
                else:
                    # Check if the specific model is available
                    models_data = response.json()
                    available_models = [
                        m["name"] for m in models_data.get("models", [])
                    ]
                    if model not in available_models:
                        print(
                            f"[DocumentService] WARNING: Model '{model}' not found in available models: {available_models}"
                        )
                    else:
                        print(f"[DocumentService] Model '{model}' is available")
        except Exception as e:
            print(
                f"[DocumentService] WARNING: Could not check Ollama service status: {e}"
            )

        # Warm up the model to reduce timeout issues
        if progress_callback:
            progress_callback("Warming up AI model...\n", 7)
        await self._warm_up_model(model)

        if progress_callback:
            progress_callback("Preparing document analysis...\n", 10)

        try:
            # Extract text from all documents
            document_texts = []
            file_paths = []

            for i, file_data in enumerate(files):
                filename = str(file_data["filename"])
                file_content = file_data["content"]

                if progress_callback:
                    progress_callback(
                        f"Processing file {i+1}/{len(files)}: {filename}\n",
                        5 + (i / len(files)) * 5,
                    )

                # Handle base64-encoded content from frontend
                if isinstance(file_content, str):
                    try:
                        file_content = base64.b64decode(file_content)
                    except Exception as e:
                        print(
                            f"[DocumentService] Failed to decode base64 for {filename}: {e}"
                        )
                        raise

                if not self.is_supported_file(filename):
                    print(f"[DocumentService] Unsupported file type: {filename}")
                    raise Exception(f"Unsupported file type: {filename}")

                file_path = await self.save_uploaded_file(file_content, filename)
                file_paths.append(file_path)

                if progress_callback:
                    progress_callback(
                        f"Extracting text from {filename}...\n",
                        10 + (i / len(files)) * 5,
                    )

                try:
                    text = await self.extract_text_from_file(file_path, filename)
                    if isinstance(text, dict) and text.get("type") == "image_based_pdf":
                        # Add progress callback to indicate detection is complete
                        if progress_callback:
                            progress_callback(
                                "Image-based PDF detected, prompting for user choice...",
                                15,
                            )
                        # Return special response for frontend to prompt user
                        return {
                            "status": "image_based_pdf",
                            "filename": filename,
                            "file_path": file_path,
                            "message": text.get("message"),
                        }
                    document_texts.append(
                        {"filename": filename, "text": text, "length": len(text)}
                    )
                    print(
                        f"[DocumentService] Extracted text from {filename}: {len(text)} chars"
                    )
                except Exception as extraction_error:
                    error_msg = str(extraction_error)
                    if "image-based" in error_msg or "no extractable text" in error_msg:
                        # Special handling for image-based PDFs
                        raise Exception(
                            f"Unable to extract text from {filename}. "
                            f"This appears to be an image-based PDF (scanned document) that requires OCR (Optical Character Recognition) to extract text. "
                            f"Please try uploading a text-based PDF or a document with selectable text."
                        )
                    else:
                        # General extraction error
                        raise Exception(
                            f"Failed to extract text from {filename}: {error_msg}. "
                            f"Please ensure the file is not corrupted and contains readable text."
                        )

            if progress_callback:
                progress_callback("Combining document content...\n", 15)

            # Check if any text was extracted
            total_extracted_text = sum(doc["length"] for doc in document_texts)
            if total_extracted_text == 0:
                raise Exception(
                    "No text could be extracted from any of the uploaded documents. "
                    "This may happen if:\n"
                    "- The documents are image-based PDFs (scanned documents)\n"
                    "- The documents are corrupted or empty\n"
                    "- The documents contain no selectable text\n\n"
                    "Please try uploading documents with selectable text or text-based PDFs."
                )

            combined_text = (
                "\n\n"
                + "=" * 50
                + "\n\n".join(
                    [
                        f"Document: {doc['filename']}\n\n{doc['text']}"
                        for doc in document_texts
                    ]
                )
            )

            should_chunk = self.should_chunk_documents(combined_text, model)
            context_length = ModelRegistry.get_context_length(model)
            estimated_tokens = ModelRegistry.estimate_tokens(combined_text)
            print(
                f"[DocumentService] Combined text length: {len(combined_text)} chars, estimated tokens: {estimated_tokens}, model context: {context_length}, should_chunk: {should_chunk}"
            )
            print(f"[DEBUG] Main decision point - should_chunk: {should_chunk}")

            if should_chunk:
                print(f"[DEBUG] Entering chunked analysis branch")
                if progress_callback:
                    progress_callback(
                        "Document is large, using chunked analysis...\n", 20
                    )
                print(f"[DocumentService] Using chunked analysis method.")
                result = await self._analyze_documents_chunked(
                    document_texts,
                    prompt,
                    model,
                    file_paths,
                    progress_callback,
                    stop_event,
                )
                result.update(
                    {
                        "chunking_info": {
                            "model_context_length": context_length,
                            "estimated_tokens": estimated_tokens,
                            "chunking_decision": "chunked",
                            "reason": f"Text too long for model context ({estimated_tokens} tokens > {int(context_length * 0.75)} available)",
                        }
                    }
                )
                print(f"[DocumentService] Chunked analysis complete. Returning result.")
                return result
            else:
                print(f"[DEBUG] Entering direct analysis branch")
                if progress_callback:
                    progress_callback(
                        "Document size is manageable, using direct analysis...\n", 20
                    )
                print(f"[DocumentService] Using direct analysis method.")
                result = await self._analyze_documents_direct(
                    combined_text,
                    prompt,
                    model,
                    file_paths,
                    progress_callback,
                    stop_event,
                )
                result.update(
                    {
                        "chunking_info": {
                            "model_context_length": context_length,
                            "estimated_tokens": estimated_tokens,
                            "chunking_decision": "direct",
                            "reason": f"Text fits in model context ({estimated_tokens} tokens <= {int(context_length * 0.75)} available)",
                        }
                    }
                )
                print(f"[DocumentService] Direct analysis complete. Returning result.")
                return result
        except Exception as e:
            print(f"[DocumentService] Exception during analysis: {e}")
            import traceback

            traceback.print_exc()
            await self._cleanup_temp_files(file_paths)
            raise e

    def _get_optimized_prompt(
        self, prompt: str, model_name: str, is_chunk: bool = False
    ) -> str:
        """Get an optimized prompt for the specific model"""

        if ModelRegistry.is_small_model(model_name):
            # For small models, use simpler, more direct prompts
            if is_chunk:
                return f"Extract key points about: {prompt}\n\nText: {{text}}\n\nKey points:"
            else:
                return f"Analyze this document and answer: {prompt}\n\nDocument:\n{{text}}\n\nAnswer:"
        else:
            # For larger models, use more detailed prompts
            if is_chunk:
                return f"""Extract key info relevant to: {prompt}

Document: {{filename}} (Chunk {{chunk_index}})
{{text}}

Brief analysis focusing on the question:"""
            else:
                return f"""Analyze these documents and answer: {prompt}

Documents:
{{text}}

Provide a clear, concise analysis addressing the question directly."""

    async def _analyze_documents_direct(
        self,
        combined_text: str,
        prompt: str,
        model: str,
        file_paths: List[str],
        progress_callback: Optional[Callable[[str, Union[int, float]], None]] = None,
        stop_event: Optional[asyncio.Event] = None,
    ) -> Dict[str, Any]:
        try:
            if progress_callback:
                progress_callback("Preparing AI model request...\n", 25)

            timeout = self.calculate_timeout(
                len(combined_text), is_chunked=False, model_name=model
            )
            print(
                f"[DocumentService] [Direct] Calling ollama_service.query_ollama with timeout={timeout}, model={model}, prompt_len={len(prompt)}, text_len={len(combined_text)}"
            )

            if progress_callback:
                progress_callback(f"Analyzing document with {model}...\n", 30)

            # Add a progress update during analysis to show it's working
            if progress_callback:
                progress_callback(f"Processing document content with {model}...\n", 50)

            # Use optimized prompt for the model
            prompt_template = self._get_optimized_prompt(prompt, model, is_chunk=False)
            full_prompt = prompt_template.format(text=combined_text)

            print(
                f"[DocumentService] [Direct] Full prompt length: {len(full_prompt)}, estimated tokens: {ModelRegistry.estimate_tokens(full_prompt)}"
            )

            # Try up to 3 times for direct analysis with increasing timeouts
            response = None
            max_attempts = (
                3 if len(combined_text) < 10000 else 2
            )  # More attempts for small documents

            for attempt in range(max_attempts):
                try:
                    # Increase timeout on each retry attempt
                    current_timeout = timeout * (
                        1.0 + (attempt * 0.5)
                    )  # 100%, 150%, 200% of original timeout
                    print(
                        f"[DocumentService] [Direct] Attempt {attempt + 1}/{max_attempts} with timeout={current_timeout}s"
                    )

                    response = await ollama_service.query_ollama(
                        full_prompt, current_timeout, model
                    )
                    print(
                        f"[DocumentService] [Direct] Model response received. Length: {len(str(response))}"
                    )
                    break
                except Exception as e:
                    if attempt < max_attempts - 1:  # Not the last attempt
                        print(
                            f"[DocumentService] [Direct] Attempt {attempt + 1} failed: {e}"
                        )
                        print(
                            f"[DocumentService] [Direct] Retrying with longer timeout..."
                        )
                        # Small delay before retry to allow model to recover
                        await asyncio.sleep(2)
                    else:
                        # Last attempt failed
                        print(
                            f"[DocumentService] [Direct] All {max_attempts} attempts failed"
                        )
                        raise e

            if progress_callback:
                progress_callback("Analysis completed successfully", 100)

            return {
                "status": "success",
                "analysis": response,
                "documents_processed": len(file_paths),
                "total_text_length": len(combined_text),
                "method": "direct",
                "timeout_used": timeout,
            }
        except Exception as e:
            print(f"[DocumentService] [Direct] Exception: {e}")
            import traceback

            traceback.print_exc()
            raise
        finally:
            await self._cleanup_temp_files(file_paths)

    async def _analyze_documents_chunked(
        self,
        document_texts: List[Dict[str, Any]],
        prompt: str,
        model: str,
        file_paths: List[str],
        progress_callback: Optional[Callable[[str, Union[int, float]], None]] = None,
        stop_event: Optional[asyncio.Event] = None,
    ) -> Dict[str, Any]:
        try:
            if progress_callback:
                progress_callback("Preparing document chunks...\n", 25)

            all_chunks = []
            for doc in document_texts:
                chunks = self.chunk_text(doc["text"], model)
                for i, chunk in enumerate(chunks):
                    all_chunks.append(
                        {"filename": doc["filename"], "chunk_index": i, "text": chunk}
                    )
            timeout_per_chunk = self.calculate_timeout(
                0, is_chunked=True, model_name=model
            )
            print(
                f"[DocumentService] [Chunked] Total chunks: {len(all_chunks)}, timeout per chunk: {timeout_per_chunk}"
            )

            if progress_callback:
                progress_callback(f"Split document into {len(all_chunks)} chunks\n", 30)

            chunk_analyses = []

            # Calculate dynamic progress range based on model and document characteristics
            avg_chunk_size = (
                sum(len(chunk["text"]) for chunk in all_chunks) // len(all_chunks)
                if all_chunks
                else 0
            )
            progress_start, progress_end = self.calculate_progress_range(
                len(all_chunks), model, avg_chunk_size
            )
            total_chunks = len(all_chunks)
            progress_per_chunk = (progress_end - progress_start) / max(total_chunks, 1)

            # Add initial progress update for large files
            if total_chunks > 10 and progress_callback:
                progress_callback(
                    f"Starting analysis of {total_chunks} document chunks...\n",
                    int(progress_start + 1),
                )
                # Small delay to ensure progress is visible
                await asyncio.sleep(0.1)

            for i, chunk in enumerate(all_chunks):
                # Check if analysis should be stopped
                if stop_event and stop_event.is_set():
                    print("[DocumentService] Analysis stopped by user request")

                    # Return partial results if we have any
                    if chunk_analyses:
                        print(
                            f"[DocumentService] Returning partial results from {len(chunk_analyses)} chunks"
                        )
                        combined_analyses = "\n\n".join(
                            [
                                f"From {analysis['filename']} (Chunk {analysis['chunk_index'] + 1}):\n{analysis['analysis']}"
                                for analysis in chunk_analyses
                            ]
                        )

                        return {
                            "status": "partial",
                            "analysis": f"Analysis was stopped by the user after processing {len(chunk_analyses)} out of {total_chunks} chunks.\n\nPartial analysis results:\n\n{combined_analyses}",
                            "documents_processed": len(document_texts),
                            "chunks_analyzed": len(chunk_analyses),
                            "total_chunks": total_chunks,
                            "method": "chunked_partial",
                            "timeout_per_chunk": timeout_per_chunk,
                        }
                    else:
                        raise Exception("Analysis stopped by user request")

                # Add debugging for stop event checking
                if stop_event:
                    print(
                        f"[DocumentService] [Chunked] Processing chunk {i+1}/{total_chunks}, stop_event.is_set(): {stop_event.is_set()}"
                    )

                if progress_callback:
                    progress = progress_start + (i * progress_per_chunk)
                    progress_callback(
                        f"Analyzing chunk {i+1}/{total_chunks} ({chunk['filename']})...\n",
                        int(progress),
                    )

                # Use optimized prompt for the model
                prompt_template = self._get_optimized_prompt(
                    prompt, model, is_chunk=True
                )
                chunk_prompt = prompt_template.format(
                    text=chunk["text"],
                    filename=chunk["filename"],
                    chunk_index=chunk["chunk_index"] + 1,
                )

                print(
                    f"[DocumentService] [Chunked] Calling ollama_service.query_ollama for chunk {chunk['chunk_index']+1} of {chunk['filename']} (chunk_len={len(chunk['text'])})"
                )

                # Add progress update during AI processing for large files
                if total_chunks > 10 and progress_callback:
                    # For files with many chunks, add intermediate progress updates
                    mid_progress = (
                        progress_start
                        + (i * progress_per_chunk)
                        + (progress_per_chunk * 0.3)
                    )
                    progress_callback(
                        f"Processing chunk {i+1}/{total_chunks} with AI...\n",
                        int(mid_progress),
                    )
                    # Small delay to ensure progress is visible
                    await asyncio.sleep(0.1)

                # Try up to 3 times for each chunk with increasing timeouts
                analysis = None
                max_chunk_attempts = (
                    3 if len(chunk["text"]) < 5000 else 2
                )  # More attempts for small chunks

                for attempt in range(max_chunk_attempts):
                    try:
                        # Increase timeout on each retry attempt
                        current_timeout = timeout_per_chunk * (
                            1.0 + (attempt * 0.5)
                        )  # 100%, 150%, 200% of original timeout
                        print(
                            f"[DocumentService] [Chunked] Chunk {chunk['chunk_index']+1} attempt {attempt + 1}/{max_chunk_attempts} with timeout={current_timeout}s"
                        )

                        analysis = await ollama_service.query_ollama(
                            chunk_prompt, current_timeout, model
                        )
                        print(
                            f"[DocumentService] [Chunked] Model response for chunk {chunk['chunk_index']+1} received. Length: {len(str(analysis))}"
                        )
                        break
                    except Exception as e:
                        if attempt < max_chunk_attempts - 1:  # Not the last attempt
                            print(
                                f"[DocumentService] [Chunked] Chunk {chunk['chunk_index']+1} attempt {attempt + 1} failed: {e}"
                            )
                            print(
                                f"[DocumentService] [Chunked] Retrying chunk {chunk['chunk_index']+1} with longer timeout..."
                            )
                            # Small delay before retry
                            await asyncio.sleep(1)
                        else:
                            # Last attempt failed - use a fallback analysis
                            print(
                                f"[DocumentService] [Chunked] All {max_chunk_attempts} attempts failed for chunk {chunk['chunk_index']+1}, using fallback"
                            )
                            analysis = f"[Analysis failed for chunk {chunk['chunk_index']+1} - content too complex for model]"

                chunk_analyses.append(
                    {
                        "filename": chunk["filename"],
                        "chunk_index": chunk["chunk_index"],
                        "analysis": analysis,
                    }
                )

                # Progress after each chunk
                if progress_callback:
                    progress = progress_start + ((i + 1) * progress_per_chunk)
                    progress_callback(
                        f"Completed chunk {i+1}/{total_chunks} ({chunk['filename']})",
                        int(progress),
                    )

            if progress_callback:
                progress_callback("Combining chunk analyses...\n", 85)

            combined_analyses = "\n\n".join(
                [
                    f"From {analysis['filename']} (Chunk {analysis['chunk_index'] + 1}):\n{analysis['analysis']}"
                    for analysis in chunk_analyses
                ]
            )
            final_timeout = self.calculate_timeout(
                len(combined_analyses), is_chunked=False, model_name=model
            )

            if progress_callback:
                progress_callback(f"Synthesizing final analysis with {model}...\n", 90)

            # Use optimized final prompt for small models
            if ModelRegistry.is_small_model(model):
                final_prompt = f"Combine this information to answer: {prompt}\n\nAnalyses:\n{combined_analyses}\n\nAnswer:"
            else:
                final_prompt = f"""Synthesize this information to answer: {prompt}

Analyses:
{combined_analyses}

Provide a coherent, comprehensive answer:"""

            print(
                f"[DocumentService] [Chunked] Calling ollama_service.query_ollama for final summary with timeout={final_timeout}, combined_analyses_len={len(combined_analyses)}"
            )

            # Try up to 3 times for the final summary with increasing timeouts
            final_analysis = None
            for attempt in range(3):
                try:
                    # Increase timeout on each retry attempt
                    current_timeout = final_timeout * (
                        1.0 + (attempt * 0.5)
                    )  # 100%, 150%, 200% of original timeout
                    print(
                        f"[DocumentService] [Chunked] Final summary attempt {attempt + 1}/3 with timeout={current_timeout}s"
                    )

                    final_analysis = await ollama_service.query_ollama(
                        final_prompt, current_timeout, model
                    )
                    print(
                        f"[DocumentService] [Chunked] Final summary response received. Length: {len(str(final_analysis))}"
                    )
                    break
                except Exception as e:
                    if attempt < 2:  # Not the last attempt
                        print(
                            f"[DocumentService] [Chunked] Final summary attempt {attempt + 1} failed: {e}"
                        )
                        print(
                            f"[DocumentService] [Chunked] Retrying with longer timeout..."
                        )
                        # Small delay before retry
                        await asyncio.sleep(1)
                    else:
                        # Last attempt failed - use a fallback summary
                        print(
                            f"[DocumentService] [Chunked] All final summary attempts failed, using fallback"
                        )
                        final_analysis = f"[Final analysis synthesis failed - the model was unable to process the combined analysis. Here are the individual chunk analyses:\n\n{combined_analyses}]"

            if progress_callback:
                progress_callback("Analysis completed successfully\n", 100)

            return {
                "status": "success",
                "analysis": final_analysis,
                "documents_processed": len(document_texts),
                "chunks_analyzed": len(all_chunks),
                "method": "chunked",
                "timeout_per_chunk": timeout_per_chunk,
                "final_timeout": final_timeout,
            }
        except Exception as e:
            print(f"[DocumentService] [Chunked] Exception: {e}")
            import traceback

            traceback.print_exc()
            raise
        finally:
            await self._cleanup_temp_files(file_paths)

    def calculate_progress_range(
        self,
        total_chunks: int,
        model_name: str,
        avg_chunk_size: int,
        has_retries: bool = True,
    ) -> tuple[int, int]:
        """
        Calculate dynamic progress range based on model and document characteristics.

        Returns:
            tuple: (progress_start, progress_end) for chunk processing
        """
        # Base progress allocation for chunk processing - AI analysis should get most of the progress
        base_chunk_progress = 60  # 60% of total progress for chunk processing (was 25%)

        # Adjust based on model speed
        if model_name in ModelRegistry.SLOW_MODELS:
            # Slow models need more time, so allocate more progress space
            chunk_progress = base_chunk_progress * 1.2  # 72%
        elif model_name in ["phi3:mini", "llama3:latest"]:
            # Fast models can use less progress space but still need significant allocation
            chunk_progress = base_chunk_progress * 0.9  # 54%
        else:
            chunk_progress = base_chunk_progress

        # Adjust based on number of chunks - for very large files, ensure minimum progress per chunk
        if total_chunks > 20:
            # Very many chunks - ensure each chunk gets at least some progress, but cap at reasonable limits
            # For very large files, we need to be more conservative with progress allocation
            if total_chunks > 100:
                # For extremely large files, use a smaller progress per chunk to stay within bounds
                min_progress_per_chunk = 0.3  # 0.3% per chunk
            elif total_chunks > 50:
                # For very large files, use moderate progress per chunk
                min_progress_per_chunk = 0.5  # 0.5% per chunk
            else:
                # For large files, use standard progress per chunk
                min_progress_per_chunk = 1.0  # 1% per chunk

            required_progress = total_chunks * min_progress_per_chunk
            # Cap the required progress to prevent exceeding 100%
            max_allowed_progress = 70  # Leave room for final synthesis (85-100%)
            required_progress = min(required_progress, max_allowed_progress)
            chunk_progress = max(chunk_progress, required_progress)
            print(
                f"[DocumentService] Large file detected ({total_chunks} chunks), using {min_progress_per_chunk}% per chunk, total progress: {chunk_progress}%"
            )
        elif total_chunks > 10:
            # Many chunks = more processing time
            chunk_progress *= 1.2
        elif total_chunks > 5:
            # Moderate number of chunks
            chunk_progress *= 1.1
        elif total_chunks <= 2:
            # Few chunks = faster processing but still significant time
            chunk_progress *= 0.8

        # Adjust based on chunk size complexity
        if avg_chunk_size > 3000:
            # Large chunks take longer
            chunk_progress *= 1.1
        elif avg_chunk_size < 1000:
            # Small chunks are faster
            chunk_progress *= 0.9

        # Account for retry attempts
        if has_retries:
            chunk_progress *= 1.05  # 5% extra for potential retries

        # Calculate start and end points - start from 30% (after file processing)
        progress_start = 10

        # Ensure we don't exceed 85% to leave room for final synthesis (85-100%)
        max_progress_end = 85
        progress_end = min(progress_start + int(chunk_progress), max_progress_end)

        # Ensure minimum progress range for very small files
        if progress_end <= progress_start:
            progress_end = progress_start + 10  # At least 10% progress

        print(
            f"[DocumentService] Progress calculation: chunks={total_chunks}, model={model_name}, avg_chunk_size={avg_chunk_size}, progress_range={progress_start}-{progress_end}, progress_per_chunk={((progress_end - progress_start) / max(total_chunks, 1)):.2f}%"
        )

        return progress_start, progress_end

    async def _cleanup_temp_files(self, file_paths: List[str]):
        """Clean up temporary files"""
        for file_path in file_paths:
            try:
                if os.path.exists(file_path):
                    await aiofiles.os.remove(file_path)
            except Exception:
                pass  # Ignore cleanup errors

    async def analyze_documents_ocr(
        self, file_path: str, filename: str, progress_callback=None
    ) -> dict:
        """Analyze a PDF using OCR extraction"""
        if not file_path or not isinstance(file_path, str):
            return {"status": "error", "message": "Invalid file_path for OCR analysis."}

        if pytesseract is None:
            raise Exception(
                "pytesseract is not installed. Please install it with: pip install pytesseract"
            )

        if progress_callback:
            progress_callback("Converting PDF to images...", 35)

        if convert_from_path is None:
            raise Exception(
                "pdf2image is not installed. Please install it with: pip install pdf2image"
            )

        try:
            images = convert_from_path(file_path)
            total_pages = len(images)

            if progress_callback:
                progress_callback(f"Converting PDF to {total_pages} images...", 40)

            text_parts = []
            for page_num, pil_image in enumerate(images):
                if progress_callback:
                    progress_per_page = 40 + (page_num / total_pages) * 20  # 40% to 60%
                    progress_callback(
                        f"Processing page {page_num + 1}/{total_pages} with OCR...",
                        int(progress_per_page),
                    )

                try:
                    ocr_text = pytesseract.image_to_string(pil_image)
                    text_parts.append(f"Page {page_num + 1} (OCR):\n{ocr_text}")
                except Exception as e:
                    text_parts.append(f"Page {page_num + 1} (OCR) failed: {e}")

            ocr_text = "\n\n".join(text_parts)

            if progress_callback:
                progress_callback("OCR extraction completed", 60)

            return {
                "status": "success",
                "method": "ocr",
                "filename": filename,
                "ocr_text": ocr_text,
            }
        except Exception as e:
            return {"status": "error", "message": f"OCR extraction failed: {str(e)}"}

    async def analyze_documents_vision(
        self,
        file_path: str,
        filename: str,
        prompt: str,
        model: str,
        progress_callback=None,
    ) -> dict:
        """Analyze a PDF as images using the vision model"""
        if not file_path or not isinstance(file_path, str):
            return {
                "status": "error",
                "message": "Invalid file_path for vision analysis.",
            }
        from app.services.image_service import image_service

        if convert_from_path is None:
            raise Exception(
                "pdf2image is not installed. Please install it with: pip install pdf2image"
            )

        if progress_callback:
            progress_callback("Converting PDF to images...", 35)

        images = convert_from_path(file_path)
        total_pages = len(images)

        if progress_callback:
            progress_callback(f"Converting PDF to {total_pages} images...", 40)

        files = []
        for i, img in enumerate(images):
            if progress_callback:
                progress_per_page = 40 + (i / total_pages) * 20  # 40% to 60%
                progress_callback(
                    f"Processing page {i + 1}/{total_pages}...", int(progress_per_page)
                )

            img_bytes = io.BytesIO()
            img.save(img_bytes, format="PNG")
            files.append(
                {
                    "filename": f"{filename}_page_{i+1}.png",
                    "content": img_bytes.getvalue(),
                }
            )

        if progress_callback:
            progress_callback("Analyzing images with vision model...", 65)

        # Use the image_service to analyze
        result = await image_service.analyze_images(files, prompt, model)

        if progress_callback:
            progress_callback("Vision analysis completed", 80)

        return {
            "status": "success",
            "method": "vision",
            "filename": filename,
            "vision_result": result,
        }


# Global service instance
document_service = DocumentService()
