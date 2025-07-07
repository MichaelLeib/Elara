import asyncio
import os
import tempfile
import base64
from typing import List, Dict, Any, Optional, Union, Callable
from pathlib import Path
import aiofiles
import aiofiles.os
from docx import Document
from docx.shared import Inches
import io
import json
import csv
from app.services.ollama_service import ollama_service
from app.config.settings import settings


class DocumentCreationService:
    """Service for creating and modifying various document types"""

    SUPPORTED_CREATION_FORMATS = {
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".csv": "text/csv",
        ".json": "application/json",
        ".xml": "application/xml",
        ".html": "text/html",
    }

    def __init__(self):
        self.max_retries = 3

    @property
    def creation_timeout(self) -> float:
        """Get creation timeout from settings, with fallback to default"""
        return getattr(settings, "document_timeout", 120.0)

    async def create_document_from_prompt(
        self,
        prompt: str,
        format_type: str = "docx",
        model: Optional[str] = None,
        base_content: Optional[str] = None,
        progress_callback: Optional[Callable[[str, Union[int, float]], None]] = None,
        stop_event: Optional[asyncio.Event] = None,
    ) -> Dict[str, Any]:
        """Create a document from a text prompt"""
        
        if progress_callback:
            progress_callback("Generating document content...", 10)

        # Build the creation prompt
        creation_prompt = self._build_creation_prompt(prompt, format_type, base_content)
        
        # Generate content using Ollama
        model = model or settings.OLLAMA_MODEL
        
        try:
            if progress_callback:
                progress_callback("Querying AI model for content generation...", 30)
                
            # Check for stop event
            if stop_event and stop_event.is_set():
                raise Exception("Document creation stopped by user request")

            content = await ollama_service.query_ollama(
                creation_prompt, 
                timeout=self.creation_timeout,
                model=model
            )
            
            if progress_callback:
                progress_callback("Content generated, creating document...", 60)

            # Create the document file
            document_path = await self._create_document_file(content, format_type, progress_callback)
            
            if progress_callback:
                progress_callback("Document created successfully!", 100)

            return {
                "status": "success",
                "file_path": document_path,
                "content": content,
                "format": format_type,
                "size": os.path.getsize(document_path),
                "filename": os.path.basename(document_path)
            }

        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "message": f"Failed to create document: {str(e)}"
            }

    async def modify_document(
        self,
        file_content: bytes,
        filename: str,
        modification_prompt: str,
        model: Optional[str] = None,
        progress_callback: Optional[Callable[[str, Union[int, float]], None]] = None,
        stop_event: Optional[asyncio.Event] = None,
    ) -> Dict[str, Any]:
        """Modify an existing document based on a prompt"""
        
        if progress_callback:
            progress_callback("Analyzing original document...", 10)

        try:
            # Save the uploaded file temporarily
            temp_path = await self._save_temp_file(file_content, filename)
            
            # Extract current content
            current_content = await self._extract_content_from_file(temp_path, filename)
            
            if progress_callback:
                progress_callback("Generating modifications...", 40)

            # Check for stop event
            if stop_event and stop_event.is_set():
                raise Exception("Document modification stopped by user request")

            # Build modification prompt
            modification_prompt_full = self._build_modification_prompt(
                current_content, modification_prompt, filename
            )
            
            # Generate modified content
            model = model or settings.OLLAMA_MODEL
            modified_content = await ollama_service.query_ollama(
                modification_prompt_full,
                timeout=self.creation_timeout,
                model=model
            )
            
            if progress_callback:
                progress_callback("Creating modified document...", 80)

            # Create the modified document
            file_ext = Path(filename).suffix.lower()
            modified_path = await self._create_document_file(modified_content, file_ext[1:], progress_callback)
            
            # Cleanup temp file
            try:
                os.remove(temp_path)
            except:
                pass

            if progress_callback:
                progress_callback("Document modified successfully!", 100)

            return {
                "status": "success",
                "file_path": modified_path,
                "original_content": current_content,
                "modified_content": modified_content,
                "format": file_ext[1:],
                "size": os.path.getsize(modified_path),
                "filename": os.path.basename(modified_path)
            }

        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "message": f"Failed to modify document: {str(e)}"
            }

    def _build_creation_prompt(self, prompt: str, format_type: str, base_content: Optional[str] = None) -> str:
        """Build a prompt for document creation"""
        
        format_instructions = {
            "docx": "Create a well-structured document with proper headings, paragraphs, and formatting.",
            "txt": "Create plain text content that is well-organized and readable.",
            "md": "Create markdown content with proper headers, lists, and formatting syntax.",
            "csv": "Create CSV data with appropriate headers and comma-separated values.",
            "json": "Create valid JSON data with proper structure and formatting.",
            "xml": "Create valid XML with proper tags and structure.",
            "html": "Create valid HTML with proper tags, structure, and semantic markup."
        }

        base_prompt = f"""Create a {format_type.upper()} document based on the following request:

{prompt}

Instructions:
- {format_instructions.get(format_type, 'Create well-structured content.')}
- Ensure the content is complete, professional, and ready to use
- Focus on clarity, accuracy, and proper organization
- Make the content comprehensive and useful

"""

        if base_content:
            base_prompt += f"""

Use this as reference or starting content:
{base_content}

"""

        base_prompt += f"""
Please provide only the {format_type.upper()} content without any additional explanation or commentary."""

        return base_prompt

    def _build_modification_prompt(self, current_content: str, modification_request: str, filename: str) -> str:
        """Build a prompt for document modification"""
        
        return f"""Modify the following document content based on the user's request:

CURRENT DOCUMENT CONTENT:
{current_content}

MODIFICATION REQUEST:
{modification_request}

Instructions:
- Make the requested changes while preserving the overall structure and format
- Maintain consistency with the existing style and tone
- Ensure all changes are integrated smoothly
- Keep any parts not mentioned in the modification request unchanged
- Provide the complete modified document content

Please provide only the modified document content without any additional explanation or commentary."""

    async def _save_temp_file(self, file_content: bytes, filename: str) -> str:
        """Save uploaded file to temporary location"""
        temp_dir = Path(tempfile.gettempdir()) / "elara_document_creation"
        temp_dir.mkdir(exist_ok=True)
        
        file_path = temp_dir / f"{os.urandom(8).hex()}_{filename}"
        
        async with aiofiles.open(file_path, "wb") as f:
            await f.write(file_content)
            
        return str(file_path)

    async def _extract_content_from_file(self, file_path: str, filename: str) -> str:
        """Extract text content from a file"""
        ext = Path(filename).suffix.lower()
        
        if ext == ".docx":
            return await self._extract_from_docx(file_path)
        elif ext in [".txt", ".md"]:
            return await self._extract_from_text(file_path)
        elif ext == ".json":
            return await self._extract_from_json(file_path)
        elif ext == ".csv":
            return await self._extract_from_csv(file_path)
        elif ext in [".html", ".xml"]:
            return await self._extract_from_text(file_path)  # Read as text
        else:
            raise Exception(f"Unsupported file format for modification: {ext}")

    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    
            return "\n\n".join(text_parts)
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")

    async def _extract_from_text(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                return await f.read()
        except UnicodeDecodeError:
            async with aiofiles.open(file_path, "r", encoding="latin-1") as f:
                return await f.read()

    async def _extract_from_json(self, file_path: str) -> str:
        """Extract formatted JSON content"""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                content = await f.read()
                # Parse and reformat for better readability
                parsed = json.loads(content)
                return json.dumps(parsed, indent=2)
        except Exception as e:
            # If parsing fails, return raw content
            return await self._extract_from_text(file_path)

    async def _extract_from_csv(self, file_path: str) -> str:
        """Extract CSV content in readable format"""
        try:
            content = await self._extract_from_text(file_path)
            # Parse CSV and format nicely
            reader = csv.reader(io.StringIO(content))
            rows = list(reader)
            
            if not rows:
                return content
                
            # Format as readable text
            formatted_rows = []
            for row in rows:
                formatted_rows.append(", ".join(row))
                
            return "\n".join(formatted_rows)
        except Exception:
            return await self._extract_from_text(file_path)

    async def _create_document_file(
        self,
        content: str,
        format_type: str,
        progress_callback: Optional[Callable[[str, Union[int, float]], None]] = None
    ) -> str:
        """Create a document file from content"""
        
        temp_dir = Path(tempfile.gettempdir()) / "elara_created_documents"
        temp_dir.mkdir(exist_ok=True)
        
        timestamp = int(asyncio.get_event_loop().time())
        filename = f"created_document_{timestamp}.{format_type}"
        file_path = temp_dir / filename
        
        if format_type == "docx":
            return await self._create_docx_file(content, str(file_path), progress_callback)
        else:
            return await self._create_text_file(content, str(file_path), format_type)

    async def _create_docx_file(
        self,
        content: str,
        file_path: str,
        progress_callback: Optional[Callable[[str, Union[int, float]], None]] = None
    ) -> str:
        """Create a DOCX file from content"""
        try:
            doc = Document()
            
            # Split content into paragraphs and process
            paragraphs = content.split('\n\n')
            
            for i, paragraph_text in enumerate(paragraphs):
                if paragraph_text.strip():
                    # Check if it looks like a heading (starts with #, all caps, etc.)
                    if self._is_heading(paragraph_text):
                        heading = doc.add_heading(paragraph_text.strip(), level=1)
                    else:
                        para = doc.add_paragraph(paragraph_text.strip())
                        
                if progress_callback and len(paragraphs) > 10:
                    progress = 70 + (i / len(paragraphs)) * 20
                    progress_callback(f"Creating DOCX structure... ({i+1}/{len(paragraphs)})", progress)
            
            doc.save(file_path)
            return file_path
            
        except Exception as e:
            raise Exception(f"Failed to create DOCX file: {str(e)}")

    async def _create_text_file(self, content: str, file_path: str, format_type: str) -> str:
        """Create a text-based file from content"""
        try:
            async with aiofiles.open(file_path, "w", encoding="utf-8") as f:
                await f.write(content)
            return file_path
        except Exception as e:
            raise Exception(f"Failed to create {format_type} file: {str(e)}")

    def _is_heading(self, text: str) -> bool:
        """Determine if text should be formatted as a heading"""
        text = text.strip()
        return (
            text.startswith('#') or  # Markdown heading
            (text.isupper() and len(text.split()) <= 5) or  # All caps short text
            text.endswith(':') or  # Ends with colon
            (len(text) <= 50 and '\n' not in text and text.replace(' ', '').isalnum())  # Short single line
        )

    async def cleanup_temp_files(self, file_paths: List[str]):
        """Clean up temporary files"""
        for file_path in file_paths:
            try:
                if os.path.exists(file_path):
                    await aiofiles.os.remove(file_path)
                    print(f"[DocumentCreationService] Cleaned up temp file: {file_path}")
            except Exception as e:
                print(f"[DocumentCreationService] Failed to cleanup {file_path}: {e}")

# Create global instance
document_creation_service = DocumentCreationService()